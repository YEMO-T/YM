import io
import os
import logging
from pptx import Presentation
from pptx.util import Inches, Pt
from docx import Document
from docx.shared import Pt as DocxPt
from typing import List, Dict, Optional

logger = logging.getLogger(__name__)

# NOTE: 模板文件存储目录，与 template_service.py 中保持一致
TEMPLATE_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'data', 'templates')

from backend.utils.pptx_builder import PPTXBuilder


def generate_pptx_stream(slides: List[Dict], template_id: Optional[str] = None) -> io.BytesIO:
    """
    将幻灯片数据转换为 PPTX 文件流。
    采用 PPTXBuilder 增强版，支持布局名称自动匹配。
    """
    template_path = None
    if template_id:
        template_path = os.path.join(TEMPLATE_DIR, f"{template_id}.pptx")
        if not os.path.exists(template_path):
            logger.warning(f"⚠ 对应的模板文件不存在: {template_path}")
            template_path = None
            
    builder = PPTXBuilder(template_path)
    for slide_data in slides:
        builder.add_slide_from_data(slide_data)
        
    return builder.build_stream()

def generate_docx_stream(lesson_plan: Dict) -> io.BytesIO:
    """
    将教案数据转换为 DOCX 文件流
    """
    doc = Document()
    
    # 标题
    title = doc.add_heading(lesson_plan.get('title', '教学教案'), 0)
    
    # 教学目标
    doc.add_heading('一、教学目标', level=1)
    objectives = lesson_plan.get('objectives', [])
    for obj in objectives:
        doc.add_paragraph(obj, style='List Bullet')
        
    # 教学过程
    doc.add_heading('二、教学过程', level=1)
    process = lesson_plan.get('process', [])
    for step in process:
        p = doc.add_paragraph()
        p.add_run(f"{step.get('stage', '环节')} ({step.get('duration', '时间')}): ").bold = True
        p.add_run(step.get('content', ''))
        
    # 课后作业
    doc.add_heading('三、课后作业', level=1)
    doc.add_paragraph(lesson_plan.get('homework', '无'))
    
    # 保存到内存流
    docx_stream = io.BytesIO()
    doc.save(docx_stream)
    docx_stream.seek(0)
    return docx_stream
